from flask import Flask, request, send_file, jsonify, render_template
from flask_cors import CORS
import os
import fitz
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
from datetime import datetime

app = Flask(__name__)
CORS(app)

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'

def extract_text_from_pdf(pdf_path, lang='eng+fas'):
    try:
        pdf_document = fitz.open(pdf_path)
        all_text = []
        total_pages = len(pdf_document)
        
        for page_num in range(total_pages):
            page = pdf_document[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img, lang=lang)
            all_text.append({
                'page': page_num + 1,
                'text': text.strip()
            })
        
        pdf_document.close()
        return all_text, total_pages
    except Exception as e:
        raise Exception(f"خطا در استخراج متن: {str(e)}")

def create_word_document(extracted_text, filename, total_pages):
    try:
        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        
        title = doc.add_heading(filename.replace('.pdf', ''), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        separator = doc.add_paragraph('─' * 60)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for page_data in extracted_text:
            page_heading = doc.add_heading(f"صفحه {page_data['page']}", level=1)
            page_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            if page_data['text']:
                paragraphs = page_data['text'].split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for run in p.runs:
                            run.font.size = Pt(12)
                            run.font.name = 'Arial'
            else:
                doc.add_paragraph('[متنی یافت نشد]')
            
            if page_data['page'] < total_pages:
                doc.add_page_break()
        
        output_filename = f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        doc.save(output_path)
        return output_path, output_filename
    except Exception as e:
        raise Exception(f"خطا در ساخت فایل Word: {str(e)}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_pdf():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'فایلی آپلود نشده است'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'فایل انتخاب نشده است'}), 400
        
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'error': 'فایل باید PDF باشد'}), 400
        
        lang = request.form.get('lang', 'eng+fas')
        
        unique_id = str(uuid.uuid4())[:8]
        original_filename = os.path.splitext(file.filename)[0]
        temp_filename = f"{unique_id}.pdf"
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_filename)
        
        file.save(temp_path)
        
        extracted_text, total_pages = extract_text_from_pdf(temp_path, lang=lang)
        output_path, output_filename = create_word_document(extracted_text, original_filename, total_pages)
        
        if os.path.exists(temp_path):
            os.remove(temp_path)
        
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.remove(temp_path)
        return jsonify({'error': f'خطا در پردازش: {str(e)}'}), 500

@app.route('/health')
def health():
    return jsonify({'status': 'OK'})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
